#!/usr/bin/env python3

import datetime as dt
import sqlite3
import zipfile
import os
import json
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.dates import DateFormatter
import io
import tempfile

import smtplib
from email.message import EmailMessage
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.chdir(os.path.dirname(os.path.abspath(__file__)))
DIR_PATH = os.path.dirname(__file__)
CONFIG_PATH = os.path.join(DIR_PATH, 'autoreport_config.json')

with open(CONFIG_PATH) as file:
    config = json.load(file)

EMAIL = config["EMAIL"]
PASSWORD = config["PASSWORD"]
MAIL_SERVER = config["MAIL_SERVER"]
DESTINATION = config["DESTINATION"]
DATABASE_PATH = config["DATABASE_PATH"]

con = sqlite3.connect(DATABASE_PATH)
cur = con.cursor()

first_day_last_month = "strftime('%s', date('now', 'start of month', '-1 month'))"
last_day_last_month = "strftime('%s', date('now', 'start of month', '-1 day'))"


res = cur.execute(f"SELECT DISTINCT sensor_id FROM sensors "
                  f"WHERE (time BETWEEN {first_day_last_month} "
                  f"AND {last_day_last_month}) ORDER BY sensor_id;")
sensors = res.fetchall()
buffer = io.BytesIO()

res = cur.execute(
    f"SELECT strftime('%m_%Y', date('now', 'start of month', '-1 month'))")
last_month = res.fetchone()[0]
last_month = dt.datetime.strptime(last_month, "%m_%Y")
doc = docx.Document()
doc.add_heading(
    f"Meranie teploty a vlhkosti "
    f"{last_month.strftime('%m/%y')}", 0)
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


for sensor in sensors:
    sensor = sensor[0]
    res = pd.read_sql(f"SELECT time, temperature, humidity FROM sensors "
                      f"WHERE sensor_id=? AND "
                      f"(time BETWEEN {first_day_last_month} AND {last_day_last_month})", con, params=[sensor])
    res["time"] = pd.to_datetime(res["time"], unit="s")
    fig, ax1 = plt.subplots()
    ax1.scatter(res["time"], res["temperature"], color="red", s=0.1)
    ax1.set_xlabel("Day", fontsize=12)
    ax1.set_ylabel("Temperature (°C)", color="red", fontsize=12)
    ax1.tick_params(axis="y", labelcolor="red")
    ax1.xaxis.set_major_formatter(DateFormatter("%d"))

    ax2 = ax1.twinx()
    ax2.scatter(res["time"], res["humidity"], color="blue", s=0.1)
    ax2.set_ylabel("Relative Humidity (%)", color="blue", fontsize=12)
    ax2.tick_params(axis="y", labelcolor="blue")

    plt.title(f"Temperature and humidity "
              f"{last_month.strftime('%m/%y')} sensor = {sensor}")
    plt.savefig(buffer, dpi=600)
    buffer.seek(0, io.SEEK_SET)
    doc.add_picture(buffer, width=Inches(5))
    buffer.seek(0, io.SEEK_SET)
    buffer.truncate(0)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

buffer.close()
report = tempfile.TemporaryFile()
doc.save(report)

res = cur.execute(f"SELECT * FROM sensors "
                  f"WHERE (time BETWEEN {first_day_last_month} AND {last_day_last_month});")
lines = res.fetchall()

raw_data = tempfile.TemporaryFile()
for line in lines:
    mp = map(float, line)
    mp = map(str, mp)
    string = ",".join(list(mp)) + "\n"
    byte = string.encode(encoding="utf-8")
    raw_data.write(byte)


zipf = tempfile.NamedTemporaryFile()
name = zipf.name
report.seek(0, io.SEEK_SET)
raw_data.seek(0, io.SEEK_SET)
with zipfile.ZipFile(zipf, "w", zipfile.ZIP_DEFLATED) as zipped:
    zipped.writestr(
        f"report_{last_month.strftime("%m_%y")}.docx", report.read())
    zipped.writestr(
        f"raw_data_{last_month.strftime("%m_%y")}.csv", raw_data.read())


report.close()
raw_data.close()

msg = EmailMessage()

msg.set_content(f"Report merania za obdobie {last_month.strftime("%m/%y")}. "
                f"Toto je automatizovaná správa, prosím neodpovedajte na ňu.")

msg['Subject'] = f"Report merania teploty a vlhkosti "
f"{last_month.strftime("%m/%y")}"
msg['From'] = EMAIL
msg['To'] = DESTINATION

with open(name, "rb") as zipf:
    data = zipf.read()
    msg.add_attachment(data, maintype="application",
                       subtype='zip', filename=f"meranie_{last_month.strftime("%m_%y")}.zip")


try:
    with smtplib.SMTP_SSL(MAIL_SERVER, 465) as smtp:
        smtp.login(EMAIL, PASSWORD)
        smtp.send_message(msg)

except Exception as err:
    print(err)
